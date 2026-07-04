"""Synthetic document generator for benchmarking.

Deterministic (fixed-seed) generator that produces .docx files at
specified paragraph counts with varied document structure:

  - headings (multiple levels)
  - body paragraphs (short, medium, long)
  - blank lines
  - references
  - numbered lists
  - bullet lists
  - nested lists
  - tables
  - mixed formatting (bold, italic, font sizes, alignment)

Each document is reproducible for regression benchmarking.
"""

from __future__ import annotations

import os
import random
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


BENCHMARK_DIR = Path(__file__).parent / "docs"
SEED = 42
RNG = random.Random(SEED)

SENTENCES_SHORT = [
    "The quick brown fox jumps over the lazy dog.",
    "Python is a high-level programming language.",
    "Document analysis remains an active research area.",
    "Natural language processing has advanced significantly.",
    "Machine learning models require large datasets.",
]

SENTENCES_MED = [
    "The results of this study demonstrate that the proposed method outperforms existing approaches across multiple evaluation metrics.",
    "In recent years, there has been growing interest in automated document understanding systems that can extract meaning from unstructured text.",
    "The analysis was conducted using a combination of statistical methods and rule-based heuristics to ensure robust classification.",
    "Table 1 summarizes the key characteristics of each document type included in the evaluation corpus.",
    "Previous work in this area has focused primarily on structural layout analysis rather than semantic understanding.",
]

SENTENCES_LONG = [
    (
        "This paper presents a comprehensive framework for document structure analysis that combines traditional "
        "layout analysis techniques with modern natural language processing methods. Our approach first extracts "
        "physical blocks from the document, then normalizes the text content, extracts linguistic features, "
        "assigns semantic roles to each block, and finally builds a hierarchical region tree that captures "
        "the document's organizational structure. We evaluate our framework on a diverse corpus of academic "
        "papers, business reports, and technical documents."
    ),
    (
        "The experimental results reveal several important findings. First, the feature extraction stage "
        "achieves high accuracy in identifying structural elements such as headings, paragraphs, and lists. "
        "Second, the classification stage demonstrates robust performance across different document types "
        "and formatting styles. Third, the region detection algorithm successfully identifies the hierarchical "
        "organization of documents, including front matter, main content sections, and back matter. "
        "These results suggest that our approach provides a solid foundation for downstream applications "
        "such as format detection and compliance validation."
    ),
    (
        "Looking ahead, there are several directions for future work. One promising avenue is extending "
        "the parser to support additional input formats such as PDF, HTML, and Markdown. Another important "
        "direction is improving the classification accuracy for edge cases such as documents with unusual "
        "formatting, mixed languages, or embedded multimedia content. Additionally, integrating the document "
        "structure analysis with content extraction techniques could enable more sophisticated applications "
        "such as automated document summarization, question answering, and knowledge base construction."
    ),
]

JOURNAL_NAMES = [
    "Journal of Document Engineering",
    "International Journal of NLP",
    "Proceedings of the ACL",
    "IEEE Transactions on Knowledge and Data Engineering",
    "ACM Computing Surveys",
]

FIRST_NAMES = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller", "Davis"]
YEARS = list(range(2019, 2026))


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else None
    if run:
        run.font.name = "Times New Roman"


def _add_body(doc: Document, sentences: list[str], count: int = 1) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(RNG.choice(sentences))
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def _add_varied_body(doc: Document) -> None:
    """Add a mix of short, medium, and long body paragraphs."""
    n = RNG.randint(2, 5)
    for _ in range(n):
        if RNG.random() < 0.3:
            _add_body(doc, SENTENCES_SHORT)
        elif RNG.random() < 0.6:
            _add_body(doc, SENTENCES_MED)
        else:
            _add_body(doc, SENTENCES_LONG)


def _add_references(doc: Document, count: int = 5) -> None:
    p = doc.add_paragraph()
    run = p.add_run("References")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i in range(count):
        ref = doc.add_paragraph()
        author = RNG.choice(FIRST_NAMES)
        year = RNG.choice(YEARS)
        journal = RNG.choice(JOURNAL_NAMES)
        ref_text = f"{author}, A. ({year}). A comprehensive study of document analysis techniques. {journal}, 15({RNG.randint(1, 4)}), {RNG.randint(100, 999)}–{RNG.randint(1000, 1099)}."
        ref.add_run(ref_text).font.size = Pt(10)
        ref.paragraph_format.left_indent = Inches(0.5)
        ref.paragraph_format.first_line_indent = Inches(-0.5)


def _add_numbered_list(doc: Document, count: int = 4) -> None:
    for i in range(count):
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {RNG.choice(SENTENCES_SHORT)}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)


def _add_bullet_list(doc: Document, count: int = 4) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        run = p.add_run(f"• {RNG.choice(SENTENCES_SHORT)}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)


def _add_nested_list(doc: Document, depth: int = 2) -> None:
    for i in range(3):
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {RNG.choice(SENTENCES_SHORT)}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Inches(0.5)
        if depth > 1:
            for j in range(2):
                sp = doc.add_paragraph()
                srun = sp.add_run(f"  a. {RNG.choice(SENTENCES_SHORT)}")
                srun.font.name = "Times New Roman"
                srun.font.size = Pt(11)
                sp.paragraph_format.left_indent = Inches(1.0)


def _add_table(doc: Document, rows: int = 4, cols: int = 3) -> None:
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"{'Header' if i == 0 else 'Cell'} {i},{j}"


PAGE_SIZE_ESTIMATE = 5  # avg paragraphs per page


def generate_docx(target_paragraphs: int, output_path: str | Path) -> None:
    """Generate a deterministic .docx file with target paragraph count.

    The document mixes headings, body text, lists, tables, and references.
    Generation is deterministic (seed=42).
    """
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("A Comprehensive Analysis")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta_lines = [
        f"Author: {RNG.choice(FIRST_NAMES)}, J.",
        "Department of Computer Science",
        "University of Technology",
        f"Date: {RNG.choice(YEARS)}-{RNG.randint(1, 12):02d}-{RNG.randint(1, 28):02d}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para_count = len(doc.paragraphs)

    section_heading_count = max(1, target_paragraphs // 50)

    # Add sections dynamically based on target paragraph count
    section_topics = [
        "Introduction",
        "Background",
        "Literature Review",
        "Methodology",
        "Experimental Setup",
        "Results",
        "Discussion",
        "Analysis",
        "Evaluation",
        "Conclusion",
        "Future Work",
    ]

    section_idx = 0

    while para_count < target_paragraphs:
        remaining = target_paragraphs - para_count

        # Decide what to add next
        if remaining < 3:
            _add_body(doc, SENTENCES_SHORT, count=1)
            para_count += 1
            continue

        choice = RNG.random()

        if choice < 0.10 and section_idx < len(section_topics):
            _add_heading(doc, section_topics[section_idx], level=1 if section_idx == 0 else RNG.randint(1, 3))
            section_idx += 1
            para_count += 1

        elif choice < 0.15 and section_idx > 1:
            _add_heading(doc, f"{RNG.choice(['Detailed', 'Additional', 'Extended'])} {RNG.choice(['Analysis', 'Results', 'Considerations'])}", level=RNG.randint(2, 3))
            para_count += 1

        elif choice < 0.60:
            _add_varied_body(doc)
            para_count += RNG.randint(2, 5)

        elif choice < 0.70:
            _add_numbered_list(doc, count=RNG.randint(2, 4))
            para_count += RNG.randint(2, 4)

        elif choice < 0.78:
            _add_bullet_list(doc, count=RNG.randint(2, 4))
            para_count += RNG.randint(2, 4)

        elif choice < 0.83 and target_paragraphs >= 100:
            _add_nested_list(doc, depth=2)
            para_count += 5

        elif choice < 0.88 and target_paragraphs >= 50:
            _add_table(doc, rows=RNG.randint(3, 6), cols=RNG.randint(2, 4))
            # Tables don't add to paragraph count directly; add a caption
            cap = doc.add_paragraph()
            cr = cap.add_run(f"Table {RNG.randint(1, 10)}. Experimental results.")
            cr.italic = True
            cr.font.size = Pt(10)
            para_count += 1

        elif choice < 0.93 and target_paragraphs >= 30:
            # Blank line
            doc.add_paragraph()
            para_count += 1

        else:
            _add_references(doc, count=RNG.randint(3, 6))
            para_count += RNG.randint(4, 7)

    # Ensure we hit the target approximately
    actual = len(doc.paragraphs)
    if actual > target_paragraphs:
        # Trim excess paragraphs (delete from end)
        p_elements = doc.element.body.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        while len(p_elements) > target_paragraphs:
            doc.element.body.remove(p_elements[-1])
            p_elements = doc.element.body.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")

    doc.save(str(output_path))


def generate_all(target_dir: str | None = None) -> dict[int, Path]:
    """Generate all benchmark documents."""
    if target_dir is None:
        target_dir = str(BENCHMARK_DIR)
    os.makedirs(target_dir, exist_ok=True)

    sizes = [10, 100, 500, 1000, 2500, 5000]
    paths = {}
    for size in sizes:
        path = Path(target_dir) / f"bench_{size}para.docx"
        if not path.exists():
            print(f"  Generating {size} paragraphs → {path}...")
            generate_docx(size, path)
        else:
            print(f"  Already exists: {path}")
        paths[size] = path
    return paths


if __name__ == "__main__":
    generate_all()
    print("Done.")
