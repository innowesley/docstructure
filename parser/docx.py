"""DOCX parser — extracts physical blocks from a .docx file.

This is the core extraction from acewriter's analyzer.py, refactored into
the DocStructure block model. The parser produces only physical blocks;
semantic classification is a separate pipeline stage.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docstructure.core.common import ParagraphRole, Provenance, Signal
from docstructure.core.document import Document
from docstructure.core.nodes import (
    BlockFeatures,
    Edge,
    ListBlock,
    ListItem,
    PageBreak,
    ParagraphBlock,
    RegionNode,
    Run,
    SectionBreak,
    TableBlock,
    TableCell,
    TableRow,
)
from docstructure.parser.base import Parser, ParserCapabilities


# ─────────────────────────────────────────────────────────────
# Constants (from acewriter analyzer.py)
# ─────────────────────────────────────────────────────────────

SECTION_NAMES: set[str] = {
    "introduction", "background", "literature review", "related work",
    "methodology", "methods", "method", "approach",
    "results", "findings", "analysis",
    "discussion",
    "conclusion", "conclusions", "summary",
    "abstract", "executive summary",
    "references", "bibliography", "works cited",
    "acknowledgements", "acknowledgments",
    "appendix", "appendices",
    "index", "glossary",
    "list of figures", "list of tables", "list of algorithms",
    "list of listings", "table of contents", "contents",
    "preface", "foreword",
}

REFERENCE_MARKERS: set[str] = {
    "references", "reference", "works cited", "bibliography",
    "works consulted", "cited works", "sources",
    "endnotes", "footnotes", "annotated bibliography",
    "selected bibliography", "selected references", "work cited",
}

APPENDIX_MARKERS: set[str] = {
    "appendix", "appendices", "supplementary material",
    "supplement", "supplements",
}

METADATA_PLACEHOLDERS: set[str] = {
    "student name", "student", "name", "author", "written by", "prepared by",
    "course", "course name", "class", "section",
    "instructor", "professor", "lecturer", "teacher", "ta", "tutor",
    "date", "submitted", "submission date", "due date",
    "assignment", "paper", "essay", "student id", "id",
    "university", "college", "department", "school",
    "semester", "term", "year",
    "institution", "faculty", "supervisor", "adviser", "advisor",
}

MONOSPACE_FONTS: set[str] = {
    "courier new", "consolas", "monaco", "lucida console",
    "courier", "monospace", "source code pro", "fira code",
    "jetbrains mono", "inconsolata",
}

BULLET_CHARS: set[str] = {"•", "●", "○", "◆", "■", "□", "▪", "▫", "–", "—"}

NUMBERED_HEADING = re.compile(r'^(\d+(?:\.\d+)*)[\.\)]\s')
ROMAN_HEADING = re.compile(r'^(X{0,3}(?:IX|IV|V?I{0,3}))[\.\)]\s', re.IGNORECASE)
CHAPTER_HEADING = re.compile(r'^chapter\s+\d+', re.IGNORECASE)


# ─────────────────────────────────────────────────────────────
# Helper: extract paragraph XML features
# ─────────────────────────────────────────────────────────────


def _has_page_break(p) -> bool:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for child in p._element:
        tag = child.tag
        if tag.endswith("}r"):
            for br in child.findall(f".//{{{ns}}}br"):
                if br.get(f"{{{ns}}}type") == "page":
                    return True
        if tag.endswith("}lastRenderedPageBreak"):
            return True
    return False


def _is_bold(p) -> bool:
    non_empty = [r for r in p.runs if r.text.strip()]
    if not non_empty:
        return False
    return all(r.bold for r in non_empty)


def _is_centered(p) -> bool:
    try:
        return p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return False


def _get_font_size(p) -> float | None:
    for r in p.runs:
        if r.font.size is not None:
            return r.font.size.pt
    return None


def _get_font_name(p) -> str:
    for r in p.runs:
        name = r.font.name
        if name:
            return name
    return ""


def _is_visual_blank(text: str, runs: list) -> bool:
    """Check if paragraph is visually blank (whitespace, NBSP, hidden, breaks)."""
    if text.strip():
        return False
    for r in runs:
        if r.font and r.font.hidden:
            continue
        t = r.text or ""
        if t.strip() and not t.strip("\u00a0"):
            return False
    return True


def _get_alignment(p) -> str:
    try:
        align = p.alignment
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        elif align == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
        elif align == WD_ALIGN_PARAGRAPH.LEFT:
            return "left"
    except Exception:
        pass
    return "left"


# ─────────────────────────────────────────────────────────────
# DOCX Parser
# ─────────────────────────────────────────────────────────────


class DOCXParser(Parser):
    """Parses a .docx file into a physical Document graph.

    Pipeline: extract paragraphs → extract runs → build ParagraphBlocks →
    detect tables → detect lists → assemble document.

    The parser NEVER classifies paragraphs semantically. It produces blocks
    with physical attributes only.
    """

    @property
    def capabilities(self) -> ParserCapabilities:
        return ParserCapabilities(
            supports_layout=False,
            supports_fonts=True,
            supports_tables=True,
            supports_images=True,
            supports_fields=True,
            supports_comments=False,
            supports_track_changes=True,
        )

    @classmethod
    def supported_mime_types(cls) -> list[str]:
        return ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

    def parse(self, source: str | bytes, **kwargs: Any) -> Document:
        doc = Document()
        doc.parser = "docx"
        doc.parser_version = "1.0"

        if isinstance(source, (str, Path)):
            docx_doc = DocxDocument(str(source))
            doc.source = str(source)
        else:
            from io import BytesIO
            docx_doc = DocxDocument(BytesIO(source))
            doc.source = "(bytes)"

        # Extract all paragraphs
        for i, p in enumerate(docx_doc.paragraphs):
            block = self._build_paragraph_block(doc, p, i)
            doc.add_node(block)

        # Extract tables
        for table in docx_doc.tables:
            table_block = self._build_table_block(doc, table)
            if table_block:
                doc.add_node(table_block)

        # Reorder blocks: interleave table blocks at their first paragraph position
        doc.nodes.sort(key=lambda n: getattr(n, "order", 0) if isinstance(n, (ParagraphBlock, PageBreak, SectionBreak)) else getattr(n, "order", 0))

        return doc

    def _build_paragraph_block(self, doc: Document, p, index: int) -> ParagraphBlock:
        text = p.text or ""
        style_name = (p.style.name or "") if p.style else ""
        font_size = _get_font_size(p)
        font_name = _get_font_name(p)

        block = ParagraphBlock(
            id=doc.next_id(),
            order=index,
            text=text,
            style_name=style_name,
            style_base=style_name,
            style_hierarchy=[style_name],
            is_visual_blank=_is_visual_blank(text, p.runs),
            is_hidden=False,
            page_break_before=_has_page_break(p),
            space_before=_get_space_before(p),
            space_after=_get_space_after(p),
        )

        # Extract style hierarchy
        if p.style and p.style.base_style:
            block.style_base = p.style.base_style.name or style_name
            bases = [block.style_base]
            s = p.style.base_style
            while s and s.base_style:
                bases.append(s.base_style.name)
                s = s.base_style
            block.style_hierarchy = [style_name] + bases

        # Outline level
        try:
            if p.paragraph_format and p.paragraph_format.outline_level is not None:
                block.outline_level = p.paragraph_format.outline_level
        except Exception:
            pass

        # Extract runs
        for ri, r in enumerate(p.runs):
            run = Run(
                id=ri,
                text=r.text or "",
                bold=bool(r.bold),
                italic=bool(r.italic),
                font_size=r.font.size.pt if r.font.size else font_size,
                font_name=r.font.name or font_name,
                is_hidden=bool(r.font.hidden) if r.font else False,
            )
            block.runs.append(run)

        # Features
        block.features = BlockFeatures(
            word_count=len(text.split()),
            sentence_count=max(1, text.count(".") + text.count("!") + text.count("?")),
            font_size=font_size,
            font_name=font_name,
            bold=_is_bold(p),
            italic=bool(block.runs and all(r.italic for r in block.runs)),
            centered=_is_centered(p),
            alignment=_get_alignment(p),
        )

        return block

    def _build_table_block(self, doc: Document, table) -> TableBlock | None:
        """Build a TableBlock from a python-docx table."""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_node = TableCell(container_id=0)
                for pi, p in enumerate(cell.paragraphs):
                    pb = self._build_paragraph_block(doc, p, doc.next_id())
                    pb.in_table = True
                    pb.container_id = 0  # Will be the table block id
                    doc.add_node(pb)
                    cell_node.child_ids.append(pb.id)
                cells.append(cell_node)
            rows.append(TableRow(cells=cells))

        if not rows:
            return None

        table_block = TableBlock(
            id=doc.next_id(),
            rows=rows,
        )
        # Set container_id on all cell blocks
        for row in rows:
            for cell in row.cells:
                cell.container_id = table_block.id

        return table_block


def _get_space_before(p) -> float | None:
    try:
        if p.paragraph_format and p.paragraph_format.space_before is not None:
            return p.paragraph_format.space_before.pt
    except Exception:
        pass
    return None


def _get_space_after(p) -> float | None:
    try:
        if p.paragraph_format and p.paragraph_format.space_after is not None:
            return p.paragraph_format.space_after.pt
    except Exception:
        pass
    return None
