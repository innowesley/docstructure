"""Generate test DOCX files for stress testing and fuzzing."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT = "/home/kunta/projects/tools/.samples/stress"


def ensure_out():
    os.makedirs(OUT, exist_ok=True)


# ── 1. APA-style paper ──────────────────────────────────────

def make_apa_paper():
    doc = Document()
    # Title page
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph("The Effects of Cognitive Load on Decision Making")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph("Jane Doe")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("University of Example")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Abstract
    p = doc.add_paragraph("Abstract")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "This study examined the relationship between cognitive load and decision-making "
        "accuracy. Participants (N = 150) completed a series of decisions under varying "
        "cognitive load conditions. Results indicate that higher cognitive load significantly "
        "reduced decision accuracy, F(2, 147) = 12.34, p < .001. These findings have "
        "implications for understanding decision-making in high-stakes environments."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_page_break()

    # Introduction
    p = doc.add_paragraph("Introduction")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Decision-making is a fundamental cognitive process that has been extensively "
        "studied in psychology and neuroscience. However, the impact of cognitive load on "
        "decision quality remains poorly understood (Smith, 2020; Johnson & Lee, 2019). "
        "The present study aims to address this gap in the literature."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph(
        "Cognitive load theory suggests that working memory has limited capacity "
        "(Sweller, 1988). When this capacity is exceeded, decision quality may suffer. "
        "Previous research has demonstrated this effect in medical diagnosis (Patel et al., "
        "2018), financial decision-making (Zhang, 2021), and consumer choice (Brown, 2020)."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Method
    p = doc.add_paragraph("Method")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph("Participants")
    r = p.runs[0]
    r.italic = True

    doc.add_paragraph(
        "A total of 150 undergraduate students (82 female, 68 male; M age = 19.8 years, "
        "SD = 1.5) participated in this study for course credit."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph("Materials and Procedure")
    r = p.runs[0]
    r.italic = True

    doc.add_paragraph(
        "Participants completed a computerized decision-making task while simultaneously "
        "performing a secondary task to manipulate cognitive load. The secondary task "
        "involved tracking a moving stimulus on screen (low load) or counting backward by "
        "threes (high load)."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Results
    p = doc.add_paragraph("Results")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "A one-way ANOVA revealed a significant main effect of cognitive load on decision "
        "accuracy, F(2, 147) = 12.34, p < .001, η² = .14. Post-hoc comparisons using "
        "Tukey's HSD indicated that accuracy was significantly lower in the high load "
        "condition (M = 62.3%) compared to both the low load (M = 78.9%) and control "
        "(M = 81.2%) conditions."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Discussion
    p = doc.add_paragraph("Discussion")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "The results support the hypothesis that cognitive load impairs decision-making "
        "accuracy. This finding is consistent with previous research and extends our "
        "understanding of the mechanisms underlying decision fatigue."
    ).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # References
    doc.add_page_break()
    p = doc.add_paragraph("References")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    refs = [
        "Brown, T. (2020). Consumer choice under pressure. Journal of Consumer Research, 47(3), 412-428.",
        "Johnson, M., & Lee, S. (2019). Decision-making in complex environments. Cognitive Psychology, 105, 1-22.",
        "Patel, V. L., Yoskowitz, N. A., & Arocha, J. F. (2018). Cognitive load in medical diagnosis. Medical Education, 52(4), 380-392.",
        "Smith, J. (2020). The psychology of decision-making. Annual Review of Psychology, 71, 201-228.",
        "Sweller, J. (1988). Cognitive load during problem solving. Cognitive Science, 12(2), 257-285.",
        "Zhang, Y. (2021). Financial decisions under cognitive load. Journal of Behavioral Finance, 22(1), 45-62.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    path = os.path.join(OUT, "apa_paper.docx")
    doc.save(path)
    return path


# ── 2. Business report ──────────────────────────────────────

def make_business_report():
    doc = Document()
    p = doc.add_paragraph("Q4 2025 Financial Performance Report")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Prepared by: Finance Department")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("January 15, 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    p = doc.add_paragraph("Executive Summary")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Revenue increased 12% year-over-year to $45.2M, driven primarily by growth in "
        "the Enterprise segment. Operating expenses were well-controlled at 68% of revenue, "
        "resulting in a net profit of $8.7M. Key initiatives for Q1 2026 include market "
        "expansion in APAC and the launch of Product X."
    )

    p = doc.add_paragraph("Revenue Analysis")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    # Table
    table = doc.add_table(rows=6, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["Segment", "Q4 2025", "Q4 2024", "Change"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Enterprise", "$22.1M", "$18.5M", "+19.5%"],
        ["SMB", "$12.8M", "$11.9M", "+7.6%"],
        ["Consumer", "$8.3M", "$8.1M", "+2.5%"],
        ["Government", "$2.0M", "$1.8M", "+11.1%"],
        ["Total", "$45.2M", "$40.3M", "+12.2%"],
    ]
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val

    doc.add_paragraph("")  # spacing

    doc.add_paragraph(
        "Enterprise segment growth was driven by three major contract renewals and two "
        "new logo acquisitions. The SMB segment saw steady growth from the self-serve "
        "channel. Consumer revenue remained flat, reflecting market saturation."
    )

    p = doc.add_paragraph("Cost Analysis")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Cost of goods sold decreased to 42% of revenue (down from 44% in Q4 2024), "
        "attributed to improved cloud infrastructure efficiency. R&D spending increased "
        "to 15% of revenue as the company invested in AI capabilities."
    )

    p = doc.add_paragraph("Outlook")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "The company expects Q1 2026 revenue in the range of $46-48M, with continued "
        "margin improvement. The APAC expansion is on track for a March launch."
    )

    path = os.path.join(OUT, "business_report.docx")
    doc.save(path)
    return path


# ── 3. Resume ───────────────────────────────────────────────

def make_resume():
    doc = Document()
    p = doc.add_paragraph("Alex Morgan")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("alex.morgan@email.com | (555) 123-4567 | linkedin.com/in/alexmorgan")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph("")  # separator

    # Summary
    p = doc.add_paragraph("Professional Summary")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph(
        "Experienced software engineer with 8+ years building distributed systems. "
        "Proven track record of leading cross-functional teams and delivering at scale."
    )

    # Experience
    p = doc.add_paragraph("Experience")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph("Senior Software Engineer | TechCorp Inc. | 2020–Present")
    r = p.runs[0]
    r.bold = True
    doc.add_paragraph("• Led migration of monolith to microservices (12 services)", style="List Bullet")
    doc.add_paragraph("• Reduced p95 latency by 40% through caching optimization", style="List Bullet")
    doc.add_paragraph("• Mentored 4 junior engineers", style="List Bullet")

    p = doc.add_paragraph("Software Engineer | StartupXYZ | 2017–2020")
    r = p.runs[0]
    r.bold = True
    doc.add_paragraph("• Built real-time analytics pipeline handling 10M events/day", style="List Bullet")
    doc.add_paragraph("• Designed REST API serving 50K requests/minute", style="List Bullet")

    # Education
    p = doc.add_paragraph("Education")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph("M.S. Computer Science | State University | 2017")
    r = p.runs[0]
    r.bold = True
    doc.add_paragraph("B.S. Computer Science | State University | 2015", style="List Bullet")

    # Skills
    p = doc.add_paragraph("Skills")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph("Languages: Python, Go, TypeScript, Rust")
    doc.add_paragraph("Tools: Kubernetes, Docker, PostgreSQL, Redis, Kafka")
    doc.add_paragraph("Cloud: AWS (certified), GCP")

    path = os.path.join(OUT, "resume.docx")
    doc.save(path)
    return path


# ── 4. Thesis chapter ────────────────────────────────────────

def make_thesis_chapter():
    doc = Document()
    p = doc.add_paragraph("Chapter 3: Methodology")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This chapter describes the methodological approach used to investigate the "
        "research questions posed in Chapter 1. The chapter is organized into four main "
        "sections: research design, participants, materials, and procedure."
    )

    p = doc.add_paragraph("3.1 Research Design")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "A mixed-methods design was employed, combining quantitative survey data with "
        "qualitative semi-structured interviews. This approach was chosen to provide both "
        "breadth and depth of understanding (Creswell & Plano Clark, 2018)."
    )

    p = doc.add_paragraph("3.1.1 Quantitative Component")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph(
        "The quantitative component consisted of a cross-sectional survey administered to "
        "500 participants. The survey included validated scales for measuring engagement, "
        "satisfaction, and intent to persist."
    )

    p = doc.add_paragraph("3.1.2 Qualitative Component")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph(
        "Semi-structured interviews were conducted with 20 participants selected from the "
        "survey respondents. Interview questions explored the experiences underlying the "
        "quantitative findings."
    )

    p = doc.add_paragraph("3.2 Participants")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Participants were recruited through university mailing lists and social media "
        "platforms. Inclusion criteria required participants to be enrolled in a "
        "degree-granting program and aged 18 or older."
    )

    # Demographics table
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["Demographic", "n", "%"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    demographics = [
        ["Gender", "", ""],
        ["  Female", "280", "56%"],
        ["  Male", "210", "42%"],
        ["  Non-binary", "10", "2%"],
    ]
    for ri, row in enumerate(demographics):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val

    doc.add_paragraph("")
    doc.add_paragraph(
        "The majority of participants were in their first year of study (42%), followed "
        "by second year (31%), third year (18%), and fourth year or beyond (9%)."
    )

    p = doc.add_paragraph("3.3 Materials")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "The survey instrument consisted of three validated scales:"
    )
    doc.add_paragraph("• Academic Engagement Scale (AES; Martin, 2007)", style="List Bullet")
    doc.add_paragraph("• Course Satisfaction Questionnaire (CSQ; Smith, 2015)", style="List Bullet")
    doc.add_paragraph("• Persistence Intent Scale (PIS; Tinto, 1993)", style="List Bullet")

    p = doc.add_paragraph("3.4 Procedure")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(
        "Data collection occurred over a six-week period during the Fall 2024 semester. "
        "Participants completed the online survey (approximately 20 minutes) and were then "
        "invited to participate in a follow-up interview."
    )

    # Equations (represented as code-like)
    doc.add_paragraph("")
    p = doc.add_paragraph("Equation 1: Sample size calculation")
    r = p.runs[0]
    r.italic = True
    doc.add_paragraph("n = (Z² × p × (1-p)) / E²")

    path = os.path.join(OUT, "thesis_chapter.docx")
    doc.save(path)
    return path


# ── 5. Weird formatting doc ──────────────────────────────────

def make_weird_formatting():
    doc = Document()
    p = doc.add_paragraph("DOCUMENT WITH WEIRD FORMATTING")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("Mixed Font Sizes and Styles")
    r = p.runs[0]
    r.font.size = Pt(8)
    r = p.add_run(" NORMAL ")
    r.font.size = Pt(12)
    r = p.add_run(" LARGE ")
    r.font.size = Pt(18)
    r.bold = True
    r = p.add_run(" tiny ")
    r.font.size = Pt(6)
    r.italic = True

    # Multiple colors
    p = doc.add_paragraph()
    colors = [
        ("Red ", RGBColor(0xFF, 0x00, 0x00)),
        ("Green ", RGBColor(0x00, 0x80, 0x00)),
        ("Blue ", RGBColor(0x00, 0x00, 0xFF)),
        ("Purple ", RGBColor(0x80, 0x00, 0x80)),
        ("Orange ", RGBColor(0xFF, 0xA5, 0x00)),
    ]
    for text, color in colors:
        r = p.add_run(text)
        r.font.color.rgb = color
        r.bold = True

    # Superscript and subscript
    p = doc.add_paragraph("Chemical formula: H")
    r = p.add_run("2")
    r.font.subscript = True
    p.add_run("O, and E=mc")
    r = p.add_run("2")
    r.font.superscript = True

    # Strikethrough and underline
    p = doc.add_paragraph()
    r = p.add_run("This text is struck through ")
    r.font.strike = True
    r = p.add_run("and this is underlined ")
    r.font.underline = True
    r = p.add_run("and this is normal.")

    # Mixed alignment on same paragraph
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Non-breaking spaces and special chars
    p = doc.add_paragraph("Non-breaking\u00a0spaces\u00a0and\u00a0tabs\u00a0everywhere")
    p = doc.add_paragraph("Smart quotes \u201cHello\u201d and \u2018World\u2019")
    p = doc.add_paragraph("Em-dash\u2014and en-dash\u2013and ellipsis\u2026")

    # Very long single word
    doc.add_paragraph("Supercalifragilisticexpialidocioussesquipedalianismantidisestablishment")

    # Tiny paragraph
    p = doc.add_paragraph("A")
    p.runs[0].font.size = Pt(72)
    p.runs[0].bold = True

    path = os.path.join(OUT, "weird_formatting.docx")
    doc.save(path)
    return path


# ── 6. Large document (100+ pages) ───────────────────────────

def make_large_document(num_paragraphs=500):
    doc = Document()
    p = doc.add_paragraph("Large Test Document")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(num_paragraphs):
        section_num = i // 50 + 1
        if i % 50 == 0:
            p = doc.add_paragraph(f"Section {section_num}")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            doc.add_paragraph(f"This is the introductory paragraph for Section {section_num}.")

        text = (
            f"This is paragraph {i + 1} of section {section_num}. It contains enough text "
            f"to represent a realistic paragraph in a large document. "
            f"The quick brown fox jumps over the lazy dog. "
            f"Researchers have studied this phenomenon extensively in recent years. "
            f"The results demonstrate a clear pattern that warrants further investigation."
        )
        doc.add_paragraph(text)

    path = os.path.join(OUT, "large_document.docx")
    doc.save(path)
    return path


# ── Fuzz documents ──────────────────────────────────────────

def make_empty_paragraphs():
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("\u00a0")
    doc.add_paragraph("Not empty")
    doc.add_paragraph("")
    doc.add_paragraph("")
    path = os.path.join(OUT, "fuzz_empty_paragraphs.docx")
    doc.save(path)
    return path


def make_nested_tables():
    doc = Document()
    # Simple table
    outer = doc.add_table(rows=2, cols=2)
    outer.style = "Light Grid Accent 1"
    outer.cell(0, 0).text = "Outer A1"
    outer.cell(0, 1).text = "Outer B1"
    outer.cell(1, 0).text = "Outer A2"
    outer.cell(1, 1).text = "Outer B2"

    # Table with merged cells (complex)
    complex_t = doc.add_table(rows=3, cols=3)
    complex_t.style = "Table Grid"
    complex_t.cell(0, 0).text = "Header 1"
    complex_t.cell(0, 1).text = "Header 2"
    complex_t.cell(0, 2).text = "Header 3"
    for i in range(3):
        complex_t.cell(1, i).text = f"Row1 Col{i+1}"
    complex_t.cell(2, 0).text = "Row2 Col1"
    # Merge last two cells
    complex_t.cell(2, 1).merge(complex_t.cell(2, 2)).text = "Merged cells"

    path = os.path.join(OUT, "fuzz_nested_tables.docx")
    doc.save(path)
    return path


def make_hidden_text():
    doc = Document()
    doc.add_paragraph("This text is visible")
    p = doc.add_paragraph()
    r = p.add_run("This text is visible too")
    r = p.add_run("This text is HIDDEN")
    r.font.hidden = True
    r = p.add_run("This text is visible again")
    doc.add_paragraph("More visible text")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    r = p.add_run("Hidden paragraph entirely")
    r.font.hidden = True
    path = os.path.join(OUT, "fuzz_hidden_text.docx")
    doc.save(path)
    return path


def make_mixed_numbering():
    doc = Document()
    doc.add_paragraph("1. Introduction")
    doc.add_paragraph("1.1 Background")
    doc.add_paragraph("1.2 Problem Statement")
    doc.add_paragraph("I. Literature Review")
    doc.add_paragraph("A. Theoretical Framework")
    doc.add_paragraph("B. Empirical Studies")
    doc.add_paragraph("II. Methodology")
    doc.add_paragraph("2.1 Research Design")
    doc.add_paragraph("2.2 Participants")
    doc.add_paragraph("Chapter 3: Results")
    doc.add_paragraph("3.1 Descriptive Statistics")
    doc.add_paragraph("IV. Discussion")
    doc.add_paragraph("References")
    doc.add_paragraph("[1] Smith, J. (2020).")
    doc.add_paragraph("[2] Jones, K. (2019).")
    path = os.path.join(OUT, "fuzz_mixed_numbering.docx")
    doc.save(path)
    return path


def make_multiple_references():
    doc = Document()
    doc.add_paragraph("Introduction to the Topic")
    doc.add_paragraph("This is a research paper with multiple reference sections.")

    doc.add_paragraph("References")
    doc.add_paragraph("Smith, J. (2020). The first study. Journal of Research, 10(2), 100-110.")
    doc.add_paragraph("Jones, K. (2019). The second study. Journal of Science, 8(4), 200-210.")

    doc.add_paragraph("Methodology")
    doc.add_paragraph("The study used a mixed-methods approach.")

    doc.add_paragraph("References")
    doc.add_paragraph("Brown, T. (2021). Methods in research. Methods Journal, 5(1), 50-60.")
    doc.add_paragraph("Lee, S. (2018). Advanced methodology. Research Methods, 12(3), 300-310.")

    doc.add_paragraph("Appendix A: Survey Instrument")
    doc.add_paragraph("The survey questions are listed below.")
    path = os.path.join(OUT, "fuzz_multiple_references.docx")
    doc.save(path)
    return path


def make_no_headings():
    doc = Document()
    doc.add_paragraph("This document has no headings at all.")
    doc.add_paragraph("It is just one paragraph after another.")
    doc.add_paragraph("The classifier should mark all as body text.")
    doc.add_paragraph("There are no section breaks.")
    doc.add_paragraph("No hierarchy whatsoever.")
    doc.add_paragraph("Just plain text.")
    doc.add_paragraph("Some more text here.")
    doc.add_paragraph("And some more.")
    doc.add_paragraph("Almost done now.")
    doc.add_paragraph("The end.")
    path = os.path.join(OUT, "fuzz_no_headings.docx")
    doc.save(path)
    return path


def make_only_tables():
    doc = Document()
    t1 = doc.add_table(rows=3, cols=3)
    t1.style = "Table Grid"
    for r in range(3):
        for c in range(3):
            t1.cell(r, c).text = f"T1 R{r+1} C{c+1}"

    doc.add_paragraph("")

    t2 = doc.add_table(rows=2, cols=4)
    t2.style = "Light Grid Accent 1"
    for r in range(2):
        for c in range(4):
            t2.cell(r, c).text = f"T2 R{r+1} C{c+1}"

    path = os.path.join(OUT, "fuzz_only_tables.docx")
    doc.save(path)
    return path


def make_corrupted_docx():
    """Create a clearly corrupted .docx by writing random bytes."""
    import random
    path = os.path.join(OUT, "fuzz_corrupted.docx")
    with open(path, "wb") as f:
        f.write(bytes(random.randint(0, 255) for _ in range(1024)))
    return path

# ── Existing samples ────────────────────────────────────────

EXISTING = [
    ("sample.docx", "/home/kunta/projects/tools/.samples/sample.docx"),
    ("samplelong.docx", "/home/kunta/projects/tools/.samples/samplelong.docx"),
    ("samplefin.docx", "/home/kunta/projects/tools/.samples/samplefin.docx"),
]


if __name__ == "__main__":
    ensure_out()
    paths = {}

    print("Generating stress test documents...")
    paths["apa_paper"] = make_apa_paper()
    print(f"  apa_paper -> {paths['apa_paper']}")
    paths["business_report"] = make_business_report()
    print(f"  business_report -> {paths['business_report']}")
    paths["resume"] = make_resume()
    print(f"  resume -> {paths['resume']}")
    paths["thesis_chapter"] = make_thesis_chapter()
    print(f"  thesis_chapter -> {paths['thesis_chapter']}")
    paths["weird_formatting"] = make_weird_formatting()
    print(f"  weird_formatting -> {paths['weird_formatting']}")

    print("\nGenerating large document (this may take a moment)...")
    paths["large_document"] = make_large_document(500)
    print(f"  large_document -> {paths['large_document']}")

    print("\nGenerating fuzz documents...")
    paths["fuzz_empty"] = make_empty_paragraphs()
    paths["fuzz_tables"] = make_nested_tables()
    paths["fuzz_hidden"] = make_hidden_text()
    paths["fuzz_numbering"] = make_mixed_numbering()
    paths["fuzz_multiref"] = make_multiple_references()
    paths["fuzz_no_headings"] = make_no_headings()
    paths["fuzz_only_tables"] = make_only_tables()
    paths["fuzz_corrupted"] = make_corrupted_docx()
    for k, v in paths.items():
        size = os.path.getsize(v)
        print(f"  {k}: {size} bytes")

    print(f"\nAll documents generated in {OUT}/")
